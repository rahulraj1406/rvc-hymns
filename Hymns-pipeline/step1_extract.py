"""
========================================================
STEP 1 — EXTRACT HYMNS FROM DOCX (CORRECTED)
========================================================
Reads your hymn book Word document and extracts each
hymn into a clean structured JSON file.

HOW TO USE:
  python3 step1_extract.py

INPUT:  hymn_book.docx
OUTPUT: hymns_raw.json
========================================================
"""

import json
import re
from docx import Document

INPUT_DOCX  = "hymns_100_to_end.docx"
OUTPUT_JSON = "hymns_raw.json"


def slugify(text):
    text = text.lower().strip()
    text = re.sub(r"[^a-z0-9\s-]", "", text)
    text = re.sub(r"\s+", "-", text)
    text = re.sub(r"-+", "-", text)
    return text.strip("-")


def is_toc_line(text):
    return bool(re.match(r"^.+\t\d+\s*$", text))


def is_hymn_header(text):
    return bool(re.match(r"^Hymn\s+\d+\s*$", text.strip()))


def get_hymn_number(text):
    m = re.match(r"^Hymn\s+(\d+)", text.strip())
    return int(m.group(1)) if m else None


def extract_toc_titles(docx_path):
    doc = Document(docx_path)
    titles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if is_toc_line(text):
            title = text.split("\t")[0].strip()
            titles.append(title)
    return titles


def extract_all_hymn_blocks(docx_path):
    doc = Document(docx_path)
    paras = [p.text for p in doc.paragraphs]

    blocks = []
    current_number = None
    current_lines  = []
    in_block       = False

    for text in paras:
        stripped = text.strip()

        if is_hymn_header(stripped):
            if current_number is not None and current_lines:
                while current_lines and not current_lines[-1].strip():
                    current_lines.pop()
                blocks.append({"number": current_number, "lines": current_lines[:]})

            current_number = get_hymn_number(stripped)
            current_lines  = []
            in_block       = True
            continue

        if in_block:
            if is_toc_line(stripped):
                in_block = False
                continue
            current_lines.append(text)

    if current_number is not None and current_lines:
        while current_lines and not current_lines[-1].strip():
            current_lines.pop()
        blocks.append({"number": current_number, "lines": current_lines[:]})

    return blocks


def group_blocks_by_hymn(blocks):
    hymn_data = {}
    for block in blocks:
        num   = block["number"]
        lines = block["lines"]
        if num not in hymn_data:
            hymn_data[num] = {"lyrics_lines": lines, "translation_lines": []}
        else:
            hymn_data[num]["translation_lines"] = lines
    return hymn_data


def clean_lyrics(lines):
    cleaned = []
    for line in lines:
        line = line.rstrip()
        line = re.sub(r"  +", " ", line)
        cleaned.append(line)
    while cleaned and not cleaned[0].strip():
        cleaned.pop(0)
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    return "\n".join(cleaned)


def clean_translation(lines):
    cleaned = []
    for line in lines:
        line = line.strip()
        line = re.sub(r"  +", " ", line)
        if line:
            cleaned.append(line)
    return " ".join(cleaned)


def build_json_objects(hymn_data, toc_titles):
    results = []
    sorted_numbers = sorted(hymn_data.keys())

    title_map = {}
    for i, num in enumerate(sorted_numbers):
        if i < len(toc_titles):
            title_map[num] = toc_titles[i]
        else:
            title_map[num] = f"Hymn {num}"

    for num in sorted_numbers:
        data  = hymn_data[num]
        title = title_map.get(num, f"Hymn {num}")

        lyrics_text = clean_lyrics(data["lyrics_lines"])
        trans_text  = clean_translation(data["translation_lines"])

        first_line = ""
        for line in lyrics_text.splitlines():
            if line.strip():
                first_line = line.strip()
                break

        obj = {
            "id":        num,
            "slug":      slugify(title),
            "number":    num,
            "title":     title,
            "firstLine": first_line,
            "category":  "general",
            "lyrics": {
                "en": lyrics_text,
                "hi": "",
                "bn": "",
                "ka": ""
            },
            "translation": {
                "en": trans_text
            }
        }
        results.append(obj)

    return results


def main():
    print(f"Reading: {INPUT_DOCX}")
    print("─" * 50)

    toc_titles = extract_toc_titles(INPUT_DOCX)
    print(f"Found {len(toc_titles)} titles in Table of Contents")

    blocks = extract_all_hymn_blocks(INPUT_DOCX)
    print(f"Found {len(blocks)} total hymn blocks")

    hymn_data = group_blocks_by_hymn(blocks)
    print(f"Found {len(hymn_data)} unique hymns")

    results = build_json_objects(hymn_data, toc_titles)

    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\n✅ Extracted {len(results)} hymns → {OUTPUT_JSON}")
    print("\nSample output:")
    print("─" * 50)
    for h in results[:3]:
        print(f"Hymn {h['number']}: {h['title']}")
        print(f"  Lyrics lines  : {len(h['lyrics']['en'].splitlines())}")
        print(f"  First line    : {h['firstLine'][:70]}")
        trans = h['translation']['en']
        print(f"  Translation   : {trans[:80]}{'...' if len(trans) > 80 else ''}")
        print()


if __name__ == "__main__":
    main()
